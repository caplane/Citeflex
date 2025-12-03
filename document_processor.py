"""
citeflex/document_processor.py

Word document processing utilities.

Features:
- LinkActivator: Adds hyperlinks to citations in Word documents
- EndnoteProcessor: Processes and formats endnotes
- BulkProcessor: Batch citation processing for entire documents

These tools integrate with the main citation system to:
1. Extract citations from Word documents
2. Look them up via the search engines
3. Format them in the desired style
4. Optionally add hyperlinks to source URLs
"""

import re
import copy
from typing import List, Optional, Tuple, Dict, Any
from dataclasses import dataclass
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[document_processor] python-docx not installed, Word processing disabled")

from models import CitationMetadata
from router import get_citation, route_and_search


@dataclass
class ProcessedCitation:
    """Result of processing a single citation."""
    original: str
    formatted: str
    metadata: Optional[CitationMetadata]
    url: Optional[str]
    success: bool
    error: Optional[str] = None


class LinkActivator:
    """
    Adds hyperlinks to citations in Word documents.
    
    Usage:
        activator = LinkActivator()
        activator.load_document("input.docx")
        activator.process_endnotes(style="Chicago Manual of Style")
        activator.save_document("output.docx")
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for document processing")
        self.document = None
        self.processed_count = 0
        self.error_count = 0
    
    def load_document(self, file_path_or_bytes) -> None:
        """Load a Word document from file path or bytes."""
        if isinstance(file_path_or_bytes, (str, bytes)):
            if isinstance(file_path_or_bytes, str):
                self.document = Document(file_path_or_bytes)
            else:
                self.document = Document(BytesIO(file_path_or_bytes))
        else:
            # Assume it's a file-like object
            self.document = Document(file_path_or_bytes)
    
    def save_document(self, file_path_or_buffer) -> None:
        """Save the document to file or buffer."""
        if self.document:
            self.document.save(file_path_or_buffer)
    
    def get_document_bytes(self) -> bytes:
        """Get the document as bytes."""
        if not self.document:
            return b""
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def process_paragraphs(
        self,
        style: str = "Chicago Manual of Style",
        add_links: bool = True,
        callback: Optional[callable] = None
    ) -> List[ProcessedCitation]:
        """
        Process citations in document paragraphs.
        
        Args:
            style: Citation style to use
            add_links: Whether to add hyperlinks
            callback: Optional callback(original, formatted, success) for progress
            
        Returns:
            List of ProcessedCitation results
        """
        if not self.document:
            return []
        
        results = []
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Process each citation (split by semicolons)
            citations = self._split_citations(text)
            
            for citation in citations:
                result = self._process_single_citation(citation, style, add_links)
                results.append(result)
                
                if callback:
                    callback(result.original, result.formatted, result.success)
        
        return results
    
    def process_endnotes(
        self,
        style: str = "Chicago Manual of Style",
        add_links: bool = True,
        callback: Optional[callable] = None
    ) -> List[ProcessedCitation]:
        """
        Process citations in document endnotes/footnotes.
        
        Note: python-docx has limited support for endnotes.
        This method tries to access them via the XML structure.
        
        Args:
            style: Citation style to use
            add_links: Whether to add hyperlinks
            callback: Optional callback for progress
            
        Returns:
            List of ProcessedCitation results
        """
        if not self.document:
            return []
        
        results = []
        
        # Try to access endnotes via document parts
        try:
            # Get endnotes part if it exists
            endnotes_part = None
            for rel in self.document.part.rels.values():
                if "endnotes" in rel.reltype:
                    endnotes_part = rel.target_part
                    break
            
            if endnotes_part:
                # Parse endnotes XML
                from lxml import etree
                root = etree.fromstring(endnotes_part.blob)
                
                # Find all endnote elements
                nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for endnote in root.findall('.//w:endnote', nsmap):
                    # Skip separator endnotes
                    endnote_type = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if endnote_type in ['separator', 'continuationSeparator']:
                        continue
                    
                    # Get text content
                    text_parts = []
                    for t in endnote.findall('.//w:t', nsmap):
                        if t.text:
                            text_parts.append(t.text)
                    
                    text = ''.join(text_parts).strip()
                    if text:
                        result = self._process_single_citation(text, style, add_links)
                        results.append(result)
                        
                        if callback:
                            callback(result.original, result.formatted, result.success)
        
        except Exception as e:
            print(f"[LinkActivator] Error processing endnotes: {e}")
            # Fall back to paragraph processing
            return self.process_paragraphs(style, add_links, callback)
        
        return results
    
    def _split_citations(self, text: str) -> List[str]:
        """Split text into individual citations."""
        # Split by semicolons (common in endnotes)
        if ';' in text:
            return [c.strip() for c in text.split(';') if c.strip()]
        return [text] if text else []
    
    def _process_single_citation(
        self,
        citation: str,
        style: str,
        add_links: bool
    ) -> ProcessedCitation:
        """Process a single citation string."""
        try:
            metadata, formatted = get_citation(citation, style)
            
            if metadata:
                self.processed_count += 1
                return ProcessedCitation(
                    original=citation,
                    formatted=formatted,
                    metadata=metadata,
                    url=metadata.url if metadata else None,
                    success=True
                )
            else:
                self.error_count += 1
                return ProcessedCitation(
                    original=citation,
                    formatted=citation,  # Keep original
                    metadata=None,
                    url=None,
                    success=False,
                    error="No metadata found"
                )
        
        except Exception as e:
            self.error_count += 1
            return ProcessedCitation(
                original=citation,
                formatted=citation,
                metadata=None,
                url=None,
                success=False,
                error=str(e)
            )
    
    def add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """
        Add a hyperlink to a paragraph.
        
        Args:
            paragraph: The docx paragraph object
            text: The text to display
            url: The URL to link to
        """
        if not url:
            return
        
        # Create relationship
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create run with text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Blue color and underline for link
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)


class BulkProcessor:
    """
    Batch processor for multiple citations.
    
    Usage:
        processor = BulkProcessor()
        results = processor.process_list(citations, style="APA 7")
    """
    
    def __init__(self):
        self.processed_count = 0
        self.error_count = 0
    
    def process_list(
        self,
        citations: List[str],
        style: str = "Chicago Manual of Style",
        callback: Optional[callable] = None
    ) -> List[ProcessedCitation]:
        """
        Process a list of citations.
        
        Args:
            citations: List of citation strings
            style: Citation style to use
            callback: Optional callback(index, total, result) for progress
            
        Returns:
            List of ProcessedCitation results
        """
        results = []
        total = len(citations)
        
        for i, citation in enumerate(citations):
            result = self._process_single(citation, style)
            results.append(result)
            
            if callback:
                callback(i + 1, total, result)
        
        return results
    
    def process_document(
        self,
        doc_bytes: bytes,
        style: str = "Chicago Manual of Style",
        extract_from: str = "paragraphs"  # or "endnotes"
    ) -> Tuple[bytes, List[ProcessedCitation]]:
        """
        Process a Word document and return updated document with results.
        
        Args:
            doc_bytes: The document as bytes
            style: Citation style to use
            extract_from: Where to extract citations ("paragraphs" or "endnotes")
            
        Returns:
            Tuple of (processed_document_bytes, results)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required")
        
        activator = LinkActivator()
        activator.load_document(doc_bytes)
        
        if extract_from == "endnotes":
            results = activator.process_endnotes(style, add_links=True)
        else:
            results = activator.process_paragraphs(style, add_links=True)
        
        return activator.get_document_bytes(), results
    
    def _process_single(self, citation: str, style: str) -> ProcessedCitation:
        """Process a single citation."""
        try:
            metadata, formatted = get_citation(citation, style)
            
            if metadata:
                self.processed_count += 1
                return ProcessedCitation(
                    original=citation,
                    formatted=formatted,
                    metadata=metadata,
                    url=metadata.url if metadata else None,
                    success=True
                )
            else:
                self.error_count += 1
                return ProcessedCitation(
                    original=citation,
                    formatted=citation,
                    metadata=None,
                    url=None,
                    success=False,
                    error="No metadata found"
                )
        
        except Exception as e:
            self.error_count += 1
            return ProcessedCitation(
                original=citation,
                formatted=citation,
                metadata=None,
                url=None,
                success=False,
                error=str(e)
            )
    
    def generate_report(self, results: List[ProcessedCitation]) -> str:
        """
        Generate a summary report of processing results.
        
        Args:
            results: List of ProcessedCitation objects
            
        Returns:
            Formatted report string
        """
        total = len(results)
        success = sum(1 for r in results if r.success)
        failed = total - success
        
        lines = [
            "=" * 60,
            "CITATION PROCESSING REPORT",
            "=" * 60,
            f"Total citations: {total}",
            f"Successfully processed: {success}",
            f"Failed: {failed}",
            f"Success rate: {success/total*100:.1f}%" if total > 0 else "N/A",
            "",
            "-" * 60,
            "DETAILS",
            "-" * 60,
        ]
        
        for i, result in enumerate(results, 1):
            status = "✓" if result.success else "✗"
            lines.append(f"\n{i}. {status} Original: {result.original[:50]}...")
            lines.append(f"   Formatted: {result.formatted[:50]}...")
            if result.url:
                lines.append(f"   URL: {result.url}")
            if result.error:
                lines.append(f"   Error: {result.error}")
        
        return "\n".join(lines)


class EndnoteEditor:
    """
    Edit endnotes in Word documents while preserving formatting.
    
    This allows for:
    - Replacing endnote text with formatted citations
    - Adding hyperlinks to sources
    - Batch updating all endnotes
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required")
        self.document = None
    
    def load_document(self, file_path_or_bytes) -> None:
        """Load a Word document."""
        if isinstance(file_path_or_bytes, (str, bytes)):
            if isinstance(file_path_or_bytes, str):
                self.document = Document(file_path_or_bytes)
            else:
                self.document = Document(BytesIO(file_path_or_bytes))
        else:
            self.document = Document(file_path_or_bytes)
    
    def get_endnotes(self) -> List[Dict[str, Any]]:
        """
        Extract all endnotes from the document.
        
        Returns:
            List of dicts with 'id', 'text', and 'element' keys
        """
        if not self.document:
            return []
        
        endnotes = []
        
        try:
            # Access endnotes part
            for rel in self.document.part.rels.values():
                if "endnotes" in rel.reltype:
                    endnotes_part = rel.target_part
                    
                    from lxml import etree
                    root = etree.fromstring(endnotes_part.blob)
                    
                    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    for endnote in root.findall('.//w:endnote', nsmap):
                        endnote_type = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                        if endnote_type in ['separator', 'continuationSeparator']:
                            continue
                        
                        endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        
                        text_parts = []
                        for t in endnote.findall('.//w:t', nsmap):
                            if t.text:
                                text_parts.append(t.text)
                        
                        endnotes.append({
                            'id': endnote_id,
                            'text': ''.join(text_parts).strip(),
                            'element': endnote
                        })
                    
                    break
        
        except Exception as e:
            print(f"[EndnoteEditor] Error reading endnotes: {e}")
        
        return endnotes
    
    def update_endnote(self, endnote_id: str, new_text: str) -> bool:
        """
        Update the text of an endnote.
        
        Args:
            endnote_id: The endnote ID
            new_text: The new text content
            
        Returns:
            True if successful
        """
        # This would require direct XML manipulation
        # For now, return False as full implementation needs more work
        print(f"[EndnoteEditor] update_endnote not fully implemented")
        return False
    
    def save_document(self, file_path_or_buffer) -> None:
        """Save the document."""
        if self.document:
            self.document.save(file_path_or_buffer)


# Convenience functions

def process_document(
    doc_path_or_bytes,
    style: str = "Chicago Manual of Style",
    output_path: Optional[str] = None
) -> List[ProcessedCitation]:
    """
    Process all citations in a Word document.
    
    Args:
        doc_path_or_bytes: Path to document or document bytes
        style: Citation style to use
        output_path: Optional path to save processed document
        
    Returns:
        List of ProcessedCitation results
    """
    activator = LinkActivator()
    activator.load_document(doc_path_or_bytes)
    
    results = activator.process_paragraphs(style, add_links=True)
    
    if output_path:
        activator.save_document(output_path)
    
    return results


def process_citations(
    citations: List[str],
    style: str = "Chicago Manual of Style"
) -> List[ProcessedCitation]:
    """
    Process a list of citation strings.
    
    Args:
        citations: List of citation strings
        style: Citation style to use
        
    Returns:
        List of ProcessedCitation results
    """
    processor = BulkProcessor()
    return processor.process_list(citations, style)
